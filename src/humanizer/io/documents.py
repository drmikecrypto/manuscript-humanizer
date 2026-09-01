from __future__ import annotations

import logging
import shutil
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from humanizer.io.pdf_layout import (
    justify_line_words,
    pdf_font_name,
    reflow_to_width_budgets,
    width_budgets_from_meta,
)

logger = logging.getLogger(__name__)

TEXT_EXTENSIONS = frozenset({".txt", ".md", ".markdown", ".tex", ".rst", ".html", ".htm"})
DOCX_EXTENSIONS = frozenset({".docx"})
PDF_EXTENSIONS = frozenset({".pdf"})
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCX_EXTENSIONS | PDF_EXTENSIONS


@dataclass
class DocumentPayload:
    """Text plus metadata needed to write back in the original format."""

    text: str
    source_path: Path | None = None
    format: str = "text"
    meta: dict[str, Any] = field(default_factory=dict)


def is_supported_path(path: Path | str) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS


def _require_docx():
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            'Word support requires: pip install -e ".[formats]" or pip install python-docx'
        ) from exc
    return Document


def _require_pymupdf():
    try:
        import fitz  # type: ignore[import-untyped]
    except ImportError as exc:
        raise ImportError(
            'PDF support requires: pip install -e ".[formats]" or pip install pymupdf'
        ) from exc
    return fitz


def _detect_line_justify(
    spans: list[dict[str, Any]],
    block_bbox: tuple[float, float, float, float],
    line_text: str,
    *,
    is_last_in_block: bool,
) -> bool:
    """Conservative heuristic: only mark lines as justified when confident."""
    if is_last_in_block:
        return False
    stripped = line_text.strip()
    char_count = len(stripped)
    word_count = len(stripped.split())
    if char_count < 40:
        return False
    if word_count < 3 and char_count < 60:
        return False

    block_width = block_bbox[2] - block_bbox[0]
    if block_width <= 0:
        return False

    if spans:
        last_bbox = spans[-1].get("bbox", block_bbox)
        fill_ratio = (float(last_bbox[2]) - block_bbox[0]) / block_width
        if fill_ratio >= 0.85:
            return True

    if len(spans) >= 3:
        gaps: list[float] = []
        for i in range(len(spans) - 1):
            prev_bbox = spans[i].get("bbox")
            next_bbox = spans[i + 1].get("bbox")
            if prev_bbox and next_bbox:
                gap = float(next_bbox[0]) - float(prev_bbox[2])
                if gap > 0:
                    gaps.append(gap)
        if gaps:
            avg_gap = sum(gaps) / len(gaps)
            if avg_gap > 3:
                variance = sum((g - avg_gap) ** 2 for g in gaps) / len(gaps)
                if variance < avg_gap * 0.5:
                    return True

    return False


def _align_pdf_lines(
    original_lines: list[str],
    humanized_text: str,
    pages_meta: list[list[dict[str, Any]]],
) -> list[str]:
    """Map humanized prose onto the original line count using width-based reflow."""
    new_lines = humanized_text.splitlines()
    if len(new_lines) == len(original_lines):
        return new_lines
    if not original_lines:
        return new_lines

    flat_meta = [meta for page in pages_meta for meta in page]
    joined = humanized_text.replace("\n", " ").strip()
    if not joined:
        return original_lines

    fontsize = 11.0
    if flat_meta:
        fontsize = float(flat_meta[0].get("size", 11))

    budgets = width_budgets_from_meta(flat_meta[: len(original_lines)], original_lines, fontsize=fontsize)
    from humanizer.io.pdf_layout import approx_measure

    measure = approx_measure(fontsize)
    result = reflow_to_width_budgets(joined, original_lines, budgets, measure)
    if result is None:
        logger.warning("PDF reflow would drop tokens; keeping original line layout")
        return original_lines
    return result


def _distribute_text(original_blocks: list[str], humanized: str) -> list[str]:
    """Map humanized prose back onto the original block count (paragraphs / pages)."""
    n = len(original_blocks)
    if n == 0:
        return []
    cleaned = humanized.strip()
    if not cleaned:
        return [""] * n
    if n == 1:
        return [cleaned]

    parts = [p.strip() for p in cleaned.split("\n\n") if p.strip()]
    if len(parts) == n:
        return parts
    if len(parts) > n:
        merged = parts[: n - 1] + ["\n\n".join(parts[n - 1 :])]
        return merged

    words = cleaned.split()
    if not words:
        return [""] * n

    weights = [max(len(block.split()), 1) for block in original_blocks]
    total = sum(weights)
    out: list[str] = []
    idx = 0
    for i, weight in enumerate(weights):
        if i == n - 1:
            out.append(" ".join(words[idx:]))
            break
        take = max(1, round(len(words) * weight / total))
        out.append(" ".join(words[idx : idx + take]))
        idx += take
    return out


def _read_text(path: Path) -> DocumentPayload:
    return DocumentPayload(
        text=path.read_text(encoding="utf-8"),
        source_path=path,
        format="text",
    )


def _read_docx(path: Path) -> DocumentPayload:
    Document = _require_docx()
    doc = Document(str(path))
    paragraphs = [para.text for para in doc.paragraphs]
    text = "\n".join(paragraphs)
    return DocumentPayload(
        text=text,
        source_path=path,
        format="docx",
        meta={"paragraphs": paragraphs},
    )


def _starts_new_paragraph(line_text: str) -> bool:
    """Heuristic for letter/sign-off lines that begin a new paragraph."""
    stripped = line_text.strip()
    if not stripped:
        return True
    lower = stripped.lower()
    if len(stripped) < 40 and (
        lower.startswith(("sincerely", "dear ", "to ", "re:", "yours ", "best ", "kind "))
        or lower in {"sincerely,", "sincerely"}
    ):
        return True
    return False


def _assign_paragraph_flags(page_lines: list[str], page_meta: list[dict[str, Any]]) -> None:
    """Detect paragraph-last lines when each PDF block is a single visual line."""
    n = len(page_meta)
    if n == 0:
        return
    for i in range(n):
        page_meta[i]["is_last_in_block"] = False
    for i in range(n - 1):
        if _starts_new_paragraph(page_lines[i + 1]):
            page_meta[i]["is_last_in_block"] = True
            continue
        cur = page_meta[i]["bbox"]
        nxt = page_meta[i + 1]["bbox"]
        line_height = max(float(cur[3]) - float(cur[1]), 8.0)
        gap = float(nxt[1]) - float(cur[3])
        indent_shift = abs(float(nxt[0]) - float(cur[0]))
        if gap > line_height * 1.8 or indent_shift > 20:
            page_meta[i]["is_last_in_block"] = True
    page_meta[n - 1]["is_last_in_block"] = True


def _paragraph_groups(page_lines: list[str], page_meta: list[dict[str, Any]]) -> list[list[int]]:
    """Group line indices into paragraphs (same breaks as _assign_paragraph_flags)."""
    n = len(page_meta)
    if n == 0:
        return []
    groups: list[list[int]] = []
    current: list[int] = [0]
    for i in range(n - 1):
        if _starts_new_paragraph(page_lines[i + 1]):
            groups.append(current)
            current = [i + 1]
            continue
        cur = page_meta[i]["bbox"]
        nxt = page_meta[i + 1]["bbox"]
        line_height = max(float(cur[3]) - float(cur[1]), 8.0)
        gap = float(nxt[1]) - float(cur[3])
        indent_shift = abs(float(nxt[0]) - float(cur[0]))
        if gap > line_height * 1.8 or indent_shift > 20:
            groups.append(current)
            current = [i + 1]
        else:
            current.append(i + 1)
    groups.append(current)
    return groups


def _propagate_paragraph_justify(page_lines: list[str], page_meta: list[dict[str, Any]]) -> None:
    """If any non-last line in a paragraph is justified, justify all non-last body lines."""
    for group in _paragraph_groups(page_lines, page_meta):
        has_justify = any(
            page_meta[i]["justify"] and not page_meta[i]["is_last_in_block"] for i in group
        )
        if not has_justify:
            continue
        for i in group:
            if page_meta[i]["is_last_in_block"]:
                continue
            if len(page_lines[i].strip()) >= 40:
                page_meta[i]["justify"] = True


def _read_pdf(path: Path) -> DocumentPayload:
    fitz = _require_pymupdf()
    doc = fitz.open(str(path))
    pages_lines: list[list[str]] = []
    pages_meta: list[list[dict[str, Any]]] = []
    try:
        for page in doc:
            page_lines: list[str] = []
            page_meta: list[dict[str, Any]] = []
            data = page.get_text("dict", sort=True)
            for block in data.get("blocks", []):
                if block.get("type") != 0:
                    continue
                block_bbox = tuple(block.get("bbox", (0, 0, 0, 0)))
                block_width = float(block_bbox[2]) - float(block_bbox[0])
                block_lines = [
                    line for line in block.get("lines", [])
                    if "".join(span.get("text", "") for span in line.get("spans", [])).strip()
                ]
                for line in block_lines:
                    spans = line.get("spans", [])
                    line_text = "".join(span.get("text", "") for span in spans)
                    if not line_text.strip():
                        continue
                    bbox = line.get("bbox")
                    first = spans[0] if spans else {}
                    size = first.get("size", 11)
                    origin = first.get("origin") or (bbox[0], bbox[3])
                    page_lines.append(line_text.rstrip())
                    page_meta.append(
                        {
                            "bbox": bbox,
                            "block_bbox": block_bbox,
                            "target_width": block_width,
                            "is_last_in_block": False,
                            "justify": False,
                            "size": size,
                            "font": first.get("font", "helv"),
                            "origin": origin,
                            "color": first.get("color"),
                            "flags": first.get("flags", 0),
                            "_spans": spans,
                            "_line_text": line_text,
                        }
                    )
            _assign_paragraph_flags(page_lines, page_meta)
            for meta in page_meta:
                meta["justify"] = _detect_line_justify(
                    meta.pop("_spans"),
                    meta["block_bbox"],  # type: ignore[arg-type]
                    meta.pop("_line_text"),
                    is_last_in_block=bool(meta["is_last_in_block"]),
                )
            _propagate_paragraph_justify(page_lines, page_meta)
            pages_lines.append(page_lines)
            pages_meta.append(page_meta)
    finally:
        doc.close()

    flat_lines = [line for page in pages_lines for line in page]
    text = "\n".join(flat_lines)
    return DocumentPayload(
        text=text,
        source_path=path,
        format="pdf",
        meta={
            "pdf_pages_lines": pages_lines,
            "pdf_pages_meta": pages_meta,
            "pdf_backup": str(path),
        },
    )


def load_document(path: Path | str) -> DocumentPayload:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)
    ext = path.suffix.lower()
    if ext in TEXT_EXTENSIONS:
        return _read_text(path)
    if ext in DOCX_EXTENSIONS:
        return _read_docx(path)
    if ext in PDF_EXTENSIONS:
        return _read_pdf(path)
    raise ValueError(
        f"Unsupported file type '{ext}'. Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
    )


def _write_text(path: Path, text: str, payload: DocumentPayload) -> None:
    path.write_text(text, encoding="utf-8")


def _write_docx(path: Path, text: str, payload: DocumentPayload) -> None:
    Document = _require_docx()
    source = payload.source_path
    original_paragraphs: list[str] = payload.meta.get("paragraphs", [])
    if not original_paragraphs and source and source.exists():
        original_paragraphs = [p.text for p in Document(str(source)).paragraphs]

    if source and source.exists() and source.suffix.lower() == ".docx":
        doc = Document(str(source))
    else:
        doc = Document()

    new_paragraphs = _distribute_text(original_paragraphs or [""], text)
    if doc.paragraphs:
        for idx, para in enumerate(doc.paragraphs):
            para.text = new_paragraphs[idx] if idx < len(new_paragraphs) else ""
        for extra in new_paragraphs[len(doc.paragraphs) :]:
            doc.add_paragraph(extra)
    else:
        for block in new_paragraphs:
            doc.add_paragraph(block)
    doc.save(str(path))


def _frozen_line_indices(text: str) -> frozenset[int]:
    """Line indices that must never be redrawn (letterhead, titles, signatures)."""
    from humanizer.rewriters.outbound import split_outbound_regions

    frozen: set[int] = set()
    for region in split_outbound_regions(text):
        if not region.frozen:
            continue
        for idx in range(region.start, region.start + len(region.lines)):
            frozen.add(idx)
    return frozenset(frozen)


def _pdf_font_name(raw: str, flags: int = 0) -> str:
    """Map PDF embedded font names to PyMuPDF base-14 names when possible."""
    return pdf_font_name(raw, flags)


def _color_int_to_rgb(color: int | None) -> tuple[float, float, float] | None:
    if color is None:
        return None
    value = int(color)
    return (
        ((value >> 16) & 255) / 255.0,
        ((value >> 8) & 255) / 255.0,
        (value & 255) / 255.0,
    )


def _insert_line_preserve_style(page: Any, text: str, meta: dict[str, Any], fitz: Any) -> None:
    """Redraw one line at the original size, color, and baseline."""
    rect = fitz.Rect(meta["bbox"])
    fontsize = float(meta.get("size", 11))
    flags = int(meta.get("flags", 0))
    fontname = _pdf_font_name(str(meta.get("font", "helv")), flags=flags)
    origin = meta.get("origin") or (rect.x0, rect.y1 - (fontsize * 0.2))
    insert_kwargs: dict[str, Any] = {}
    rgb = _color_int_to_rgb(meta.get("color"))
    if rgb is not None:
        insert_kwargs["color"] = rgb

    justify = bool(meta.get("justify")) and not bool(meta.get("is_last_in_block"))
    words = text.split()
    if justify and len(words) >= 1:
        block_bbox = meta.get("block_bbox")
        target_width = float(meta.get("target_width") or (rect.x1 - rect.x0))
        x_start = float(block_bbox[0]) if block_bbox else float(rect.x0)
        y = float(origin[1]) if isinstance(origin, (tuple, list)) else float(origin.y)
        justify_line_words(
            page,
            fitz,
            words,
            x_start=x_start,
            target_width=target_width,
            y=y,
            fontsize=fontsize,
            fontname=fontname,
            insert_kwargs=insert_kwargs,
        )
        return

    kwargs: dict[str, Any] = {
        "fontsize": fontsize,
        "fontname": fontname,
        **insert_kwargs,
    }
    if isinstance(origin, (tuple, list)):
        page.insert_text(origin, text, **kwargs)
    else:
        page.insert_text(origin, text, **kwargs)


def _redact_rect_for_line(meta: dict[str, Any], bbox: Any, page_rect: Any, fitz: Any) -> Any:
    """Build redaction rectangle; use full block width for justified body lines."""
    block_bbox = meta.get("block_bbox")
    if meta.get("justify") and block_bbox:
        rect = fitz.Rect(block_bbox)
    else:
        rect = fitz.Rect(bbox)
    rect.x0 = max(page_rect.x0, rect.x0 - 1)
    if meta.get("justify") and block_bbox:
        rect.x1 = min(page_rect.x1, float(block_bbox[2]) + 1)
    else:
        rect.x1 = min(page_rect.x1, page_rect.x1 - 24)
    rect.y0 -= 1
    rect.y1 += 1
    return rect


def _write_pdf(path: Path, text: str, payload: DocumentPayload) -> None:
    """Preserve original PDF design; replace editable body lines in-place only."""
    fitz = _require_pymupdf()
    backup = payload.meta.get("pdf_backup")
    layout_path = Path(backup) if backup and Path(backup).exists() else payload.source_path
    if layout_path is None or not layout_path.exists():
        raise FileNotFoundError("Original PDF layout reference is missing; cannot preserve formatting.")

    pages_lines: list[list[str]] = payload.meta.get("pdf_pages_lines") or []
    pages_meta: list[list[dict[str, Any]]] = payload.meta.get("pdf_pages_meta") or []
    flat_old = [line for page in pages_lines for line in page]
    flat_new = _align_pdf_lines(flat_old, text, pages_meta) if flat_old else text.splitlines()
    if len(flat_new) < len(flat_old):
        flat_new.extend([""] * (len(flat_old) - len(flat_new)))

    baseline_text = "\n".join(flat_old)
    aligned_text = "\n".join(flat_new)
    from humanizer.validators.fidelity import build_quality_report

    qc_report = build_quality_report(baseline_text, aligned_text)
    if not qc_report.passed and not payload.meta.get("force_save"):
        logger.warning("PDF write QC failed; reverting to original lines (%s)", qc_report.issues)
        flat_new = list(flat_old)

    frozen = _frozen_line_indices(baseline_text)

    line_map: list[tuple[int, dict[str, Any]]] = []
    for page_num, page_lines in enumerate(pages_lines):
        for line_idx in range(len(page_lines)):
            line_map.append((page_num, pages_meta[page_num][line_idx]))

    page_edits: dict[int, list[tuple[Any, str, dict[str, Any]]]] = {}
    for line_idx, (old_line, new_line) in enumerate(zip(flat_old, flat_new)):
        if line_idx in frozen:
            continue
        if old_line.strip() == new_line.strip():
            continue
        page_num, meta = line_map[line_idx]
        page_edits.setdefault(page_num, []).append((meta["bbox"], new_line.rstrip(), meta))

    if not page_edits:
        shutil.copy2(layout_path, path)
        return

    doc = fitz.open(str(layout_path))
    try:
        for page_num, edits in page_edits.items():
            page = doc[page_num]
            edits.sort(key=lambda item: item[0][1], reverse=True)
            page_rect = page.rect
            for bbox, _new_text, meta in edits:
                rect = _redact_rect_for_line(meta, bbox, page_rect, fitz)
                page.add_redact_annot(rect, fill=(1, 1, 1))
            page.apply_redactions()
            edits.sort(key=lambda item: item[0][1])
            for _bbox, new_text, meta in edits:
                _insert_line_preserve_style(page, new_text, meta, fitz)
        doc.save(str(path), garbage=4, deflate=True)
    finally:
        doc.close()


def save_document(path: Path | str, text: str, payload: DocumentPayload) -> Path:
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)
    ext = path.suffix.lower()

    if ext in TEXT_EXTENSIONS or payload.format == "text":
        _write_text(path, text, payload)
    elif ext in DOCX_EXTENSIONS or payload.format == "docx":
        _write_docx(path, text, payload)
    elif ext in PDF_EXTENSIONS or payload.format == "pdf":
        _write_pdf(path, text, payload)
    else:
        _write_text(path.with_suffix(".txt"), text, payload)
        path = path.with_suffix(".txt")
    return path


def resolve_output_path(
    input_path: Path,
    *,
    output: Path | None = None,
    in_place: bool = False,
) -> Path:
    if in_place:
        return input_path
    if output is not None:
        return output
    return input_path.with_name(f"{input_path.stem}_humanized{input_path.suffix or '.txt'}")


def backup_file(path: Path) -> Path | None:
    if not path.exists():
        return None
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S")
    backup = path.with_name(f"{path.stem}.bak.{stamp}{path.suffix}")
    shutil.copy2(path, backup)
    return backup
